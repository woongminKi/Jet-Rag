"""W5 DE-67 — DocxParser 단위 테스트.

검증 범위
- can_parse: .docx 확장자 매칭
- heading sticky propagate (style.name 'Heading' 패턴)
- inline 텍스트 패턴 fallback (제 N 조 등)
- 표 처리 (행/셀 join + 빈 표 skip)
- 한국어 unicode 자연 처리
- iter_inner_content 의 paragraph + table 순서 보존
- corrupted DOCX → RuntimeError wrap

E2 4차 ship — 실 DOCX 자산 fixture 회귀 추가 (`DocxParserRealAssetTest`).
메모리 합성 binary 로는 잡히지 않는 실 파일 회귀를 보호. 자산 미존재 시 자동 skip.

자산 디렉토리 우선순위 (5단계, `test_pymupdf_heading.py` 패턴과 정합)
- 1순위: `<repo>/assets/public/` — 모든 컴퓨터·CI 자동 회귀
- 2순위: `<repo>/assets/` 직속 (사용자 PC raw 자료)
- 3순위: `<repo>/` 루트 직속 (다른 컴퓨터 패턴)
- 4순위: `JETRAG_TEST_DOCX_DIR` ENV 폴백
- 5단계: 자산 부재 시 자동 skip

stdlib unittest + python-docx 로 합성 — 외부 자료 의존 0.
"""

from __future__ import annotations

import io
import os
import unittest
from pathlib import Path

import docx as python_docx

os.environ.setdefault("HF_API_TOKEN", "dummy-test-token")


# repo root 자동 인식: api/tests/test_*.py → parents[2] = repo root
_REPO_ROOT = Path(__file__).resolve().parents[2]
_PUBLIC_DOCX_DIR = _REPO_ROOT / "assets" / "public"

# 공개 fixture (현재 0건 — follow-up 후보, 라이센스 검토 통과 시 추가)
_PUBLIC_DOCX_FILES: list[str] = []

# 비공개 자료 (사용자 PC `assets/` 직속, `.gitignore` 로 다른 컴퓨터엔 부재)
_PRIVATE_DOCX_FILES = [
    "승인글 템플릿1.docx",
    "승인글 템플릿3.docx",
]

# 회귀 가능 자산 = 공개 + 비공개 (5단계 우선순위로 자동 해석)
_DOCX_FILES = _PUBLIC_DOCX_FILES + _PRIVATE_DOCX_FILES


def _docx_path(name: str) -> Path:
    """5단계 우선순위로 DOCX fixture 경로 해석. 부재 시 부재 path 반환 (호출부 skipTest).

    1) `<repo>/assets/public/<name>` — 공개 fixture
    2) `<repo>/assets/<name>` — 사용자 PC raw 자료
    3) `<repo>/<name>` — 다른 컴퓨터 repo 루트 직속
    4) `$JETRAG_TEST_DOCX_DIR/<name>` — ENV 폴백
    5) 부재 → public path 반환 (exists() False)
    """
    public = _PUBLIC_DOCX_DIR / name
    if public.exists():
        return public

    assets_direct = _REPO_ROOT / "assets" / name
    if assets_direct.exists():
        return assets_direct

    repo_root_direct = _REPO_ROOT / name
    if repo_root_direct.exists():
        return repo_root_direct

    env_base = os.environ.get("JETRAG_TEST_DOCX_DIR")
    if env_base:
        env_path = Path(env_base) / name
        if env_path.exists():
            return env_path

    return public  # exists() False — 호출부에서 skipTest


def _build_docx(*, title: str | None = None, body: list[tuple[str, str]] | None = None,
                table: list[list[str]] | None = None) -> bytes:
    """합성 DOCX 생성 — body 는 (style_name, text) 리스트, table 은 행렬."""
    doc = python_docx.Document()
    if title:
        doc.add_heading(title, level=1)
    for style_name, text in (body or []):
        if style_name.startswith("Heading"):
            level = int(style_name[-1]) if style_name[-1].isdigit() else 1
            doc.add_heading(text, level=level)
        else:
            p = doc.add_paragraph(text)
            if style_name and style_name != "Normal":
                try:
                    p.style = doc.styles[style_name]
                except KeyError:
                    pass
    if table:
        rows = len(table)
        cols = len(table[0]) if rows else 0
        if rows and cols:
            t = doc.add_table(rows=rows, cols=cols)
            for r, row in enumerate(table):
                for c, val in enumerate(row):
                    t.cell(r, c).text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class CanParseTest(unittest.TestCase):
    def test_docx_extension(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        self.assertTrue(DocxParser().can_parse("file.docx", None))
        self.assertTrue(DocxParser().can_parse("FILE.DOCX", None))

    def test_non_docx_rejected(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        self.assertFalse(DocxParser().can_parse("file.pdf", None))
        self.assertFalse(DocxParser().can_parse("file.hwpx", None))
        self.assertFalse(DocxParser().can_parse("file.txt", None))


class HeadingStickyPropagateTest(unittest.TestCase):
    def test_heading_propagates_to_next_paragraphs(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        data = _build_docx(
            title="1단원 개요",
            body=[
                ("Normal", "이것은 본문 단락이다."),
                ("Heading 2", "1.1 세부 항목"),
                ("Normal", "두 번째 본문이다."),
            ],
        )
        result = DocxParser().parse(data, file_name="test.docx")
        # heading 자체도 sections 에 포함
        self.assertEqual(len(result.sections), 4)
        # 첫 heading 까지 propagate
        self.assertEqual(result.sections[0].section_title, "1단원 개요")
        self.assertEqual(result.sections[1].section_title, "1단원 개요")
        # 두 번째 heading 만나면 갱신
        self.assertEqual(result.sections[2].section_title, "1.1 세부 항목")
        self.assertEqual(result.sections[3].section_title, "1.1 세부 항목")

    def test_no_heading_keeps_title_none(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        data = _build_docx(body=[
            ("Normal", "본문 1"),
            ("Normal", "본문 2"),
        ])
        result = DocxParser().parse(data, file_name="test.docx")
        for s in result.sections:
            self.assertIsNone(s.section_title)

    def test_inline_text_pattern_fallback(self) -> None:
        """style.name 이 'Normal' 이지만 텍스트가 outline 패턴이면 heading 으로 인식."""
        from app.adapters.impl.docx_parser import _is_heading_paragraph

        self.assertTrue(_is_heading_paragraph("제 1 조 (목적)", "Normal"))
        self.assertTrue(_is_heading_paragraph("부칙", "Normal"))
        self.assertTrue(_is_heading_paragraph("별표 1", "Normal"))
        self.assertFalse(_is_heading_paragraph("일반 본문이다.", "Normal"))


class TableExtractionTest(unittest.TestCase):
    def test_table_text_joined_with_separator(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        data = _build_docx(
            title="표 예시",
            table=[["항목 A", "항목 B"], ["값 1", "값 2"]],
        )
        result = DocxParser().parse(data, file_name="test.docx")
        # 표는 별도 ExtractedSection — 마지막에 위치
        table_section = result.sections[-1]
        self.assertIn("항목 A | 항목 B", table_section.text)
        self.assertIn("값 1 | 값 2", table_section.text)

    def test_empty_table_skipped(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        # 빈 셀만 있는 표
        data = _build_docx(
            body=[("Normal", "본문")],
            table=[["", ""], ["", ""]],
        )
        result = DocxParser().parse(data, file_name="test.docx")
        # 본문 단락 1건만 (빈 표는 skip)
        text_only = [s for s in result.sections if s.text]
        self.assertEqual(len(text_only), 1)

    def test_table_inherits_current_title(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        data = _build_docx(
            title="1단원",
            body=[("Normal", "본문")],
            table=[["a", "b"], ["c", "d"]],
        )
        result = DocxParser().parse(data, file_name="test.docx")
        # 표 직전 heading "1단원" 이 표 section_title 로 propagate
        table_section = next(s for s in result.sections if "|" in s.text)
        self.assertEqual(table_section.section_title, "1단원")


class CorruptedDocxTest(unittest.TestCase):
    def test_invalid_bytes_raises_runtime_error(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        with self.assertRaises(RuntimeError) as ctx:
            DocxParser().parse(b"not a docx", file_name="bad.docx")
        self.assertIn("DOCX 파서 초기화 실패", str(ctx.exception))


class RawTextTest(unittest.TestCase):
    def test_raw_text_concat_includes_paragraphs_and_tables(self) -> None:
        from app.adapters.impl.docx_parser import DocxParser

        data = _build_docx(
            title="요약",
            body=[("Normal", "본문 단락이다.")],
            table=[["A", "B"]],
        )
        result = DocxParser().parse(data, file_name="test.docx")
        # raw_text 에 heading + 본문 + 표 모두 포함
        self.assertIn("요약", result.raw_text)
        self.assertIn("본문 단락이다.", result.raw_text)
        self.assertIn("A | B", result.raw_text)


class DocxParserRealAssetTest(unittest.TestCase):
    """E2 4차 ship — 실 DOCX 자산에 대한 회귀 보호.

    메모리 합성 (`_build_docx`) 으로는 잡히지 않는 케이스:
    - 실제 워드프로세서로 생성된 스타일 변형
    - 한국어 글꼴·복합 단락·승인 양식 표 등 외부 도구 산출물
    - 실 사용자 자산의 인코딩·메타데이터 정합

    자산 부재 시 자동 skip (CI 호환). 사용자 PC `assets/` 직속에 자료 있으면 자동 진입.
    """

    def setUp(self) -> None:
        self._target_name: str | None = None
        self._target_path: Path | None = None
        for name in _DOCX_FILES:
            path = _docx_path(name)
            if path.exists():
                self._target_name = name
                self._target_path = path
                return
        self.skipTest("실 DOCX 자산 부재 — public/private 모두 미존재")

    def test_can_parse_real_docx(self) -> None:
        """실 DOCX 파싱 — section 1건 이상 + 첫 section text non-empty."""
        from app.adapters.impl.docx_parser import DocxParser

        assert self._target_path is not None and self._target_name is not None
        data = self._target_path.read_bytes()
        result = DocxParser().parse(data, file_name=self._target_name)

        self.assertGreaterEqual(
            len(result.sections), 1,
            f"실 DOCX 파싱 시 section 1건 이상 — got {len(result.sections)}",
        )
        # text 가 1건 이상 있는 section 이 존재 (heading-only 단락만 있는 경우 회피)
        non_empty = [s for s in result.sections if s.text.strip()]
        self.assertGreaterEqual(
            len(non_empty), 1,
            "non-empty text section 0건 — 파서가 텍스트 회수 실패",
        )


if __name__ == "__main__":
    unittest.main()
